"""
app/translate/document_builder.py
===================================
The inverse of document_parser.py — reconstructs a translated document
in the SAME file format it was uploaded as (CONCEPT.md — "download it in
the same format they uploaded it in").

PDF is still not a full round-trip: document_parser.py's heading
detection (font-size heuristic for PDF, paragraph style for DOCX) is a
real structural signal, not full layout/positioning — no columns,
images, or exact spacing are captured. The generated PDF/DOCX apply
that one signal (heading vs. body) but are not a facsimile of the
source. XLSX/CSV are proper reconstructions since we control that
structure directly (the cell grid).
"""

from __future__ import annotations

import csv
import io
import logging

import openpyxl
from docx import Document as DocxDocument
from fpdf import FPDF

logger = logging.getLogger("thought_translate.document_builder")

# Devanagari-capable font candidates, in lookup order. The macOS entries
# are for local dev only; NOTO_LINUX_PATH is what's actually available in
# the Docker image once `fonts-noto-core` is installed there (see
# Dockerfile) — production depends on that path existing, not the mac ones.
_FONT_CANDIDATES = [
    "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf",
    "/System/Library/Fonts/Supplemental/DevanagariMT.ttc",
    "/System/Library/Fonts/Supplemental/Devanagari Sangam MN.ttc",
    "/System/Library/Fonts/Supplemental/ITFDevanagari.ttc",
]


class NoUnicodeFontAvailable(Exception):
    pass


def _find_font() -> str:
    import os

    for path in _FONT_CANDIDATES:
        if os.path.exists(path):
            return path
    raise NoUnicodeFontAvailable(
        "No Devanagari-capable font found on this server, install fonts-noto-core (or similar) to enable PDF export."
    )


def build_docx(paragraphs: list[str], is_heading: list[bool] | None = None) -> bytes:
    doc = DocxDocument()
    headings = is_heading or [False] * len(paragraphs)
    for para, heading in zip(paragraphs, headings):
        if heading:
            doc.add_paragraph(para, style="Heading 2")
        else:
            doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_xlsx(rows: list[list[str]]) -> bytes:
    wb = openpyxl.Workbook()
    sheet = wb.active
    for row in rows:
        sheet.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_csv(rows: list[list[str]]) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    for row in rows:
        writer.writerow(row)
    return buf.getvalue().encode("utf-8-sig")


def build_pdf(paragraphs: list[str], is_heading: list[bool] | None = None) -> bytes:
    """Plain text-flow PDF — see module docstring on why this isn't a
    layout-preserving reconstruction. Headings render larger/bold,
    everything else at body size."""
    font_path = _find_font()
    headings = is_heading or [False] * len(paragraphs)
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("Deva", "", font_path)
    for para, heading in zip(paragraphs, headings):
        pdf.set_font("Deva", size=16 if heading else 12)
        pdf.multi_cell(0, 9 if heading else 8, para)
        pdf.ln(5 if heading else 4)
    return bytes(pdf.output())
