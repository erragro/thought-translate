"""
app/translate/document_parser.py
==================================
Extraction for the accepted upload types (CONCEPT.md §1, expanded
2026-08-03 to add structure + OCR).

Two shapes come out of this module, matching §1's split:
  - ProseDocument  — PDF / Word / images: a list of paragraphs (prose),
    with a parallel `is_heading` flag per paragraph (heading detection:
    font-size heuristic for PDF via PyMuPDF, paragraph style for DOCX —
    both real signals, not guessed).
  - TableDocument  — Excel / CSV: a 2D grid of cells (structure/formulas
    untouched — only cell text gets translated, by the caller).

PDF handling now uses PyMuPDF (fitz) instead of pdfminer — it exposes
per-span font size (needed for heading detection) and can render a page
to an image (needed for OCR), which pdfminer's plain-text API couldn't
do. A page with no extractable text is treated as scanned and OCR'd via
Tesseract (pytesseract) rather than silently returning nothing — this is
the "OCR scanner" + "transcribe images" capability. Combined eng+hin
OCR regardless of declared source_lang, since Tesseract's multi-language
mode handles mixed content fine and our language scope is exactly those
two right now.
"""

from __future__ import annotations

import csv
import io
import statistics
from dataclasses import dataclass, field

import fitz  # PyMuPDF
import openpyxl
import pytesseract
from docx import Document as DocxDocument
from PIL import Image

SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".csv", ".jpg", ".jpeg", ".png"}

# Combined language mode — our scope is exactly English + Hindi right now.
_OCR_LANGS = "eng+hin"

# A text block's average font size needs to be at least this many times
# the document's median body-text size to count as a heading. Heuristic,
# not a real style/outline read (PDF doesn't reliably expose one).
_HEADING_SIZE_RATIO = 1.15


class UnsupportedFileType(Exception):
    pass


@dataclass
class ProseDocument:
    kind: str = "prose"
    paragraphs: list[str] = field(default_factory=list)
    is_heading: list[bool] = field(default_factory=list)


@dataclass
class TableDocument:
    kind: str = "table"
    rows: list[list[str]] = field(default_factory=list)


def _ocr_image_bytes(data: bytes) -> str:
    image = Image.open(io.BytesIO(data))
    return pytesseract.image_to_string(image, lang=_OCR_LANGS)


def _split_paragraphs(text: str) -> list[str]:
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    # OCR output sometimes has no blank-line breaks at all — fall back to
    # single newlines rather than returning one giant blob.
    if len(paragraphs) <= 1 and text.strip():
        paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    return paragraphs


def _extract_pdf(data: bytes) -> ProseDocument:
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        page_blocks: list[list[tuple[str, float]]] = []
        all_sizes: list[float] = []

        for page in doc:
            text_dict = page.get_text("dict")
            blocks: list[tuple[str, float]] = []
            for block in text_dict.get("blocks", []):
                if block.get("type") != 0:  # 0 = text block; skip images etc.
                    continue
                parts: list[str] = []
                sizes: list[float] = []
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        t = span.get("text", "")
                        if t.strip():
                            parts.append(t)
                            sizes.append(span.get("size", 0.0))
                block_text = " ".join(parts).strip()
                if block_text:
                    avg_size = sum(sizes) / len(sizes) if sizes else 0.0
                    blocks.append((block_text, avg_size))
                    all_sizes.append(avg_size)
            page_blocks.append(blocks)

        median_size = statistics.median(all_sizes) if all_sizes else 0.0
        heading_threshold = median_size * _HEADING_SIZE_RATIO if median_size else 0.0

        paragraphs: list[str] = []
        is_heading: list[bool] = []
        for page_index, blocks in enumerate(page_blocks):
            if blocks:
                for text, size in blocks:
                    paragraphs.append(text)
                    is_heading.append(bool(heading_threshold) and size >= heading_threshold)
            else:
                # No extractable text on this page — scanned/image content.
                # OCR it rather than silently dropping the page.
                pix = doc[page_index].get_pixmap(dpi=200)
                ocr_text = _ocr_image_bytes(pix.tobytes("png"))
                for para in _split_paragraphs(ocr_text):
                    paragraphs.append(para)
                    is_heading.append(False)

        return ProseDocument(paragraphs=paragraphs, is_heading=is_heading)
    finally:
        doc.close()


def _extract_docx(data: bytes) -> ProseDocument:
    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as exc:
        # python-docx only reads the modern .docx (zip/XML) format — a
        # real legacy .doc (binary) file lands here and needs a clear
        # message, not a confusing zip-parsing traceback.
        raise UnsupportedFileType(
            "Could not read this Word file. Legacy .doc (pre-2007) format "
            "isn't supported yet, please save as .docx and re-upload."
        ) from exc

    paragraphs: list[str] = []
    is_heading: list[bool] = []
    for p in doc.paragraphs:
        if p.text.strip():
            paragraphs.append(p.text.strip())
            style_name = (p.style.name or "").lower() if p.style else ""
            is_heading.append(style_name.startswith("heading") or style_name.startswith("title"))
    return ProseDocument(paragraphs=paragraphs, is_heading=is_heading)


def _extract_xlsx(data: bytes) -> TableDocument:
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    except Exception as exc:
        # Same class of problem as .doc above — openpyxl only reads the
        # modern .xlsx (zip/XML) format, not legacy binary .xls.
        raise UnsupportedFileType(
            "Could not read this Excel file. Legacy .xls (pre-2007) format "
            "isn't supported yet, please save as .xlsx and re-upload."
        ) from exc
    sheet = wb.active
    rows: list[list[str]] = []
    for row in sheet.iter_rows(values_only=True):
        rows.append(["" if cell is None else str(cell) for cell in row])
    return TableDocument(rows=rows)


def _extract_csv(data: bytes) -> TableDocument:
    text = data.decode("utf-8-sig", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = [list(row) for row in reader]
    return TableDocument(rows=rows)


def _extract_image(data: bytes) -> ProseDocument:
    text = _ocr_image_bytes(data)
    paragraphs = _split_paragraphs(text)
    return ProseDocument(paragraphs=paragraphs, is_heading=[False] * len(paragraphs))


def extract(filename: str, data: bytes) -> ProseDocument | TableDocument:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf(data)
    if lower.endswith(".docx") or lower.endswith(".doc"):
        return _extract_docx(data)
    if lower.endswith(".xlsx") or lower.endswith(".xls"):
        return _extract_xlsx(data)
    if lower.endswith(".csv"):
        return _extract_csv(data)
    if lower.endswith(".jpg") or lower.endswith(".jpeg") or lower.endswith(".png"):
        return _extract_image(data)
    raise UnsupportedFileType(f"Unsupported file type: {filename}")
